import os
import re
import streamlit as st
import tempfile
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import google.generativeai as genai
import markdown2
from bs4 import BeautifulSoup

st.set_page_config(page_title="Panchayat Audio Report Generator", layout="wide")
st.title("Panchayat Audio Report Generator")
st.markdown(
    "Upload multiple audio files (WAV, MP3, OGG, FLAC, M4A, MP4) to transcribe, translate to Malayalam, and generate professional Panchayat reports in DOCX format. After processing, you will be able to download your reports directly."
)

# Configure Gemini API key from secrets (REQUIRED for Gemini API)
if "gemini" in st.secrets and "api_key" in st.secrets["gemini"]:
    genai.configure(api_key=st.secrets["gemini"]["api_key"])
else:
    st.error("Gemini API key not found in secrets. Please add [gemini] api_key to your .streamlit/secrets.toml.")
    st.stop()

# Set up Google Application Credentials for other Google APIs (OPTIONAL, only if you use them)
def set_google_credentials_from_secrets():
    try:
        credentials_dict = dict(st.secrets["google"])
        with tempfile.NamedTemporaryFile(delete=False, suffix=".json", mode="w") as tmp_file:
            json.dump(credentials_dict, tmp_file)
            tmp_file_path = tmp_file.name
        os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = tmp_file_path
        st.session_state.credentials_path = tmp_file_path
        return True
    except Exception as e:
        st.error(f"Error loading credentials from Secrets: {e}")
        return False

if "credentials_set" not in st.session_state:
    st.session_state.credentials_set = set_google_credentials_from_secrets()

def cleanup():
    if "credentials_path" in st.session_state:
        try:
            os.unlink(st.session_state.credentials_path)
        except Exception:
            pass

import atexit
atexit.register(cleanup)

def get_mimetype(ext):
    ext = ext.lower()
    if ext == ".mp3":
        return "audio/mp3"
    elif ext == ".ogg":
        return "audio/ogg"
    elif ext == ".flac":
        return "audio/flac"
    elif ext == ".m4a" or ext == ".mp4":
        return "audio/mp4"
    else:
        return "audio/wav"

def improved_gemini_prompt(mal_text):
    return f"""
ഈ transcript-ന്റെ അടിസ്ഥാനത്തിൽ, താഴെ പറയുന്ന ഘടനയും നിർദേശങ്ങളും കർശനമായി പാലിച്ചുകൊണ്ട് ഒരു ഗ്രാമപഞ്ചായത്ത് വിശദമായ, ഔദ്യോഗിക ഭാഷയിൽ എഴുതിയ, പ്രൊഫഷണൽ വിശകലന റിപ്പോർട്ട് markdown format-ൽ തയ്യാറാക്കുക.

**If any section lacks sufficient information in the transcript, create representative, plausible, and professional content for that section, using your general knowledge of Panchayat reports in Kerala.**

# ഗ്രാമപഞ്ചായത്തിന്റെ യഥാർത്ഥ പേര്, ജില്ല (audio/transcript-ൽ ലഭ്യമായെങ്കിൽ മാത്രം; placeholder ഉപയോഗിക്കരുത്.)

## മോദി സർക്കാരിന്റെ നേട്ടങ്ങൾ
- transcript-ൽ നിന്നുള്ള പ്രധാന നേട്ടങ്ങൾ, പദ്ധതികൾ, അവയുടെ ഫലപ്രാപ്തി, ഉദാഹരണങ്ങൾ
- subheadings (ഉപശീർഷികങ്ങൾ) transcript-ൽ നിന്നുണ്ടെങ്കിൽ bold smaller headings ആയി ഉപയോഗിക്കുക; ഇവയില്ലെങ്കിൽ, സാധാരണയായി കാണുന്ന ഉപശീർഷികങ്ങൾ ചേർക്കുക.
- body government report style-ൽ, വിശദീകരണം, context, statistics, examples എന്നിവയോടെ എഴുതുക.

## ആരോപണങ്ങൾ / അടിസ്ഥാന പ്രശ്നങ്ങൾ
- transcript-ൽ നിന്നുള്ള പ്രധാന പ്രശ്നങ്ങൾ, ഉദാഹരണങ്ങൾ, സ്ഥലങ്ങൾ, പരിഹാരങ്ങൾ
- subheadings transcript-ൽ നിന്ന് കണ്ടെത്തി bold smaller headings ആയി എഴുതുക; ഇവയില്ലെങ്കിൽ, സാധാരണ പ്രശ്നങ്ങൾ ചേർക്കുക.
- body content paragraph-ൽ വിശദമായ വിവരണം, background, example, implication ഉൾപ്പെടുത്തുക.

## വികസന രേഖ/ദർശനം
- transcript-ൽ നിന്നുള്ള വികസന ലക്ഷ്യങ്ങൾ, പദ്ധതികൾ, നിർദേശങ്ങൾ, പുതിയ ആശയങ്ങൾ
- subheadings transcript-ൽ നിന്ന് bold smaller headings ആയി ഉപയോഗിച്ച്, ഇവയില്ലെങ്കിൽ സാധാരണ vision headings ചേർക്കുക.
- body content professional, visionary, policy-oriented language-ൽ paragraph-ൽ വിശദമായി എഴുതുക.

- Section headings bold, left aligned; subheadings bold, left aligned, smaller font; body justified alignment, bullet points/numbered lists ആവശ്യമായിടത്ത് മാത്രം ഉപയോഗിക്കുക.
- Report-ൽ ഒരു ഭാഗത്തും [Insert ... here], "---", “audio does not provide”, “audio indicates”, “audio mentions”, “not mentioned”, “unavailable”, “lack of data”, “audio-യിൽ” എന്നൊന്നും എഴുതരുത്.

Transcript:
{mal_text}
"""

def extract_panchayat_name(md_text):
    match = re.search(r'^#\s*(.+)', md_text, re.MULTILINE)
    if match:
        return match.group(1).strip()
    return "ഗ്രാമപഞ്ചായത്ത് ഡോക്യുമെന്റ്"

def markdown_to_docx(md_text):
    html = markdown2.markdown(md_text, extras=["tables"])
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()

    panchayat_name = extract_panchayat_name(md_text)
    title = doc.add_heading(panchayat_name, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    style = doc.styles['Normal']
    style.font.name = 'Noto Sans Malayalam'
    style.font.size = Pt(12)

    body = soup.body if soup.body else soup
    if not body or not list(body.children):
        return None

    for el in body.children:
        name = getattr(el, "name", None)
        if name and name.startswith("h"):
            if name == "h1":
                continue
            level = int(name[1])
            doc.add_heading(el.text, min(level, 4))
        elif name == "ul":
            for li in el.find_all("li", recursive=False):
                doc.add_paragraph(li.text, style="List Bullet")
        elif name == "ol":
            for li in el.find_all("li", recursive=False):
                doc.add_paragraph(li.text, style="List Number")
        elif name == "p":
            doc.add_paragraph(el.text)
        elif name == "table":
            rows = el.find_all("tr")
            if not rows:
                continue
            cols = rows[0].find_all(["td", "th"])
            table = doc.add_table(rows=len(rows), cols=len(cols))
            table.style = 'Table Grid'
            for i, row in enumerate(rows):
                for j, cell in enumerate(row.find_all(["td", "th"])):
                    table.cell(i, j).text = cell.text
        elif name is None and el.string and el.string.strip():
            doc.add_paragraph(el.string.strip())
    # Save the docx to a temporary file and return the bytes
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        doc.save(tmp_file.name)
        tmp_file.seek(0)
        docx_bytes = tmp_file.read()
    os.unlink(tmp_file.name)
    return docx_bytes

def transcribe_and_translate(model, audio_bytes, mimetype, fname):
    transcript_response = model.generate_content([
        {"mime_type": mimetype, "data": audio_bytes},
        "Transcribe this audio. Output only the transcript."
    ])
    transcript = transcript_response.text.strip()
    if not transcript or transcript.strip().lower() in (
        "none", "no speech detected", "could not transcribe"):
        return ""
    translation_response = model.generate_content([
        f"ഇത് മലയാളത്തിലേക്ക് വിവർത്തനം ചെയ്യുക:\n{transcript}"
    ])
    mal_text = translation_response.text.strip()
    return mal_text

def generate_professional_document(model, mal_text, fname):
    doc_prompt = improved_gemini_prompt(mal_text)
    doc_response = model.generate_content([doc_prompt])
    professional_doc_md = doc_response.text.strip()
    return professional_doc_md

if st.session_state.credentials_set:
    st.header("Step 1: Upload Audio Files")
    audio_files = st.file_uploader(
        "Upload audio files (WAV, MP3, OGG, FLAC, M4A, MP4)",
        type=["wav", "mp3", "ogg", "flac", "m4a", "mp4"],
        accept_multiple_files=True
    )

    # Initialize session state for processed files
    if "processed_files" not in st.session_state:
        st.session_state.processed_files = dict()

    if audio_files:
        model = genai.GenerativeModel("models/gemini-1.5-flash-latest")

        for audio_file in audio_files:
            fname = audio_file.name
            ext = os.path.splitext(fname)[1]
            mimetype = get_mimetype(ext)

            # Use session state to avoid re-processing
            # Use key as (filename, file_size) to handle files with same name but different contents
            file_key = (fname, audio_file.size)
            if file_key not in st.session_state.processed_files:
                with st.spinner(f"Processing: {fname}"):
                    try:
                        audio_bytes = audio_file.read()
                        mal_text = transcribe_and_translate(model, audio_bytes, mimetype, fname)
                        if not mal_text.strip():
                            st.session_state.processed_files[file_key] = None
                            st.error(f"Skipping {fname}: no transcript or translation extracted.")
                            continue
                        professional_doc_md = generate_professional_document(model, mal_text, fname)
                        if not professional_doc_md.strip():
                            st.session_state.processed_files[file_key] = None
                            st.error(f"Skipping {fname}: markdown document was not generated.")
                            continue
                        docx_bytes = markdown_to_docx(professional_doc_md)
                        if docx_bytes:
                            st.session_state.processed_files[file_key] = docx_bytes
                        else:
                            st.session_state.processed_files[file_key] = None
                            st.error(f"Failed to generate DOCX for {fname}.")
                    except Exception as e:
                        st.session_state.processed_files[file_key] = None
                        st.error(f"Error processing {fname}: {e}")

        # Now, offer download buttons for processed files (cached)
        for audio_file in audio_files:
            fname = audio_file.name
            file_key = (fname, audio_file.size)
            docx_bytes = st.session_state.processed_files.get(file_key)
            if docx_bytes:
                st.download_button(
                    label=f"Download {fname} DOCX",
                    data=docx_bytes,
                    file_name=f"{os.path.splitext(fname)[0]}_panchayat_document.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                st.success(f"DOCX ready for download: {fname}")
